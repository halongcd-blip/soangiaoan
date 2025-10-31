import streamlit as st
import time
from docx import Document
from io import BytesIO
import re # Cần để làm sạch Markdown
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

# -----------------------------------------------------------------
# 1. CẤU HÌNH "BỘ NÃO" AI
# -----------------------------------------------------------------
import google.generativeai as genai

# LẤY API KEY TỪ STREAMLIT SECRETS
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except:
    # Dùng API Key giả cho môi trường giả lập, bạn cần thay bằng API Key thật
    API_KEY = "FAKE_API_KEY_FOR_DEMO" 

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel(model_name="gemini-2.5-flash")

# -----------------------------------------------------------------
# **PROMPT GỐC (ĐÃ ĐƯỢC CẬP NHẬT THEO YÊU CẦU MỚI CỦA BẠN)**
# -----------------------------------------------------------------
PROMPT_GOC = """
CẢNH BÁO QUAN TRỌNG: TUYỆT ĐỐI KHÔNG SỬ DỤNG BẤT CỨ THẺ HTML NÀO (ví dụ: <br/>, <strong>). Hãy dùng định dạng MARKDOWN thuần túy (dấu * hoặc - cho gạch đầu dòng và xuống dòng tự động).

Bạn là một chuyên gia giáo dục Tiểu học hàng đầu Việt Nam, am hiểu sâu sắc Chương trình GDPT 2018 và kỹ thuật thiết kế Kế hoạch Bài Dạy (giáo án) theo Công văn 2345.

Nhiệm vụ của bạn là soạn một Kế hoạch bài dạy chi tiết, sáng tạo, tập trung vào phát triển năng lực và phẩm chất.

**YÊU CẦU ĐẶC BIỆT VỀ HÌNH ẢNH (NẾU CÓ):**
Nếu người dùng tải lên 1 hoặc 2 ảnh, bạn PHẢI:
1.  Phân tích hình ảnh (đây là ảnh chụp bài tập trong SGK).
2.  Trích xuất (chuyển ảnh thành chữ) **CHỈ NỘI DUNG BÀI TẬP** (không lấy hình ảnh/chi tiết thừa).
3.  Đưa nội dung bài tập (ĐỀ BÀI) này vào cột "Hoạt động của giáo viên" (trong Hoạt động 3).
4.  Soạn thảo ĐÁP ÁN/HƯỚNG DẪN GIẢI cho bài tập đó và đưa vào cột "Hoạt động của học sinh".
5.  ĐỀ BÀI và ĐÁP ÁN phải nằm trên **CÙNG MỘT HÀNG NGANG** của bảng, không được lệch dòng.

DỮ LIỆU ĐẦU VÀO:
1.  **Môn học:** {mon_hoc}
2.  **Lớp:** {lop}
3.  **Bộ sách:** {bo_sach}
4.  **Tên bài học/Chủ đề:** {ten_bai}
5.  **Yêu cầu cần đạt (Lấy từ Chương trình môn học):** {yeu_cau}
6.  **Yêu cầu tạo phiếu bài tập:** {yeu_cau_phieu}
7.  **Yêu cầu tạo sơ đồ tư duy:** {yeu_cau_mindmap} 

YÊU CẦU VỀ ĐỊNH DẠNG:
Bạn PHẢI tuân thủ tuyệt đối cấu trúc và các yêu cầu sau:

**I. Yêu cầu cần đạt**
(Phát biểu cụ thể học sinh thực hiện được việc gì; vận dụng được những gì; phẩm chất, năng lực gì.)
1.  **Về kiến thức:** (Bám sát {yeu_cau})
2.  **Về năng lực:** (Năng lực chung: Tự chủ và tự học, Giao tiếp và hợp tác, Giải quyết vấn đề và sáng tạo; Năng lực đặc thù của môn {mon_hoc})
3.  **Về phẩm chất:** (Chọn 1-2 trong 5 phẩm chất: Yêu nước, Nhân ái, Chăm chỉ, Trung thực, Trách nhiệm)

**II. Đồ dùng dạy học**
(Nêu các thiết bị, học liệu được sử dụng trong bài dạy. Nếu Yêu cầu tạo phiếu bài tập là CÓ, phải nhắc đến Phiếu bài tập trong mục này.)
1.  **Chuẩn bị của giáo viên (GV):** (Tranh ảnh, video, phiếu học tập, link game...)
2.  **Chuẩn bị của học sinh (HS):** (SGK, Vở bài tập, bút màu...)

**III. Các hoạt động dạy học chủ yếu**
**QUY TẮC QUAN TRỌNG VỀ NỘI DUNG:** (TỔNG THỜI GIAN TIẾT HỌC LÀ 35 PHÚT).
Phần này PHẢI được soạn thật **kỹ lưỡng, chi tiết và tỉ mỉ** cho từng hoạt động (a, b, c...).
Ưu tiên sử dụng các phương pháp và kỹ thuật dạy học tích cực (ví dụ: KWL, Mảnh ghép, Khăn trải bàn, Góc học tập, Trạm học tập, Đóng vai, Sơ đồ tư duy...) để phát huy tối đa năng lực và phẩm chất của học sinh theo Chương trình GDPT 2018.

**QUY TẮC QUAN TRỌNG VỀ BẢNG BIỂU:** Toàn bộ nội dung của mục 3 này PHẢI được trình bày trong **MỘT BẢNG MARKDOWN DUY NHẤT** có 2 cột.

| Hoạt động của giáo viên | Hoạt động của học sinh |
| :--- | :--- |
| **1. Hoạt động Mở đầu (Khởi động, Kết nối) (Khoảng 3-5 phút)** | |
| (Viết chi tiết các bước tổ chức, dẫn dắt vào bài, dùng dấu gạch đầu dòng `*` cho mỗi bước) | (Viết chi tiết các hoạt động tương tác, chuẩn bị của HS, dùng dấu gạch đầu dòng `*`) |
| **2. Hoạt động Hình thành kiến thức mới (Trải nghiệm, Khám phá) (Khoảng 10-12 phút)** | |
| (Viết chi tiết các bước tổ chức HS trải nghiệm, khám phá, hình thành kiến thức, dùng dấu gạch đầu dòng `*`) | (Viết chi tiết các bước HS quan sát, thảo luận, ghi chép, dùng dấu gạch đầu dòng `*`) |
| **3. Hoạt động Luyện tập, Thực hành (Khoảng 15-18 phút)** | |
| (Viết chi tiết các bước tổ chức HS áp dụng kiến thức. Nếu có ảnh, chèn ĐỀ BÀI (đã trích xuất) vào đây. Nếu có phiếu, giao phiếu ở đây. Dùng `*`) | (Viết chi tiết các bước HS thực hành. Nếu có ảnh, chèn ĐÁP ÁN/HƯỚNG DẪN GIẢI vào đây ở ô CÙNG HÀNG. Dùng `*`) |
| **4. Hoạt động Vận dụng, Trải nghiệm (Củng cố) (Khoảng 3-5 phút)** | |
| (Viết chi tiết các bước tổ chức HS liên hệ thực tế, củng cố bài, dùng dấu gạch đầu dòng `*`) | (Viết chi tiết các bước HS trả lời, cam kết hành động, dùng dấu gạch đầu dòng `*`) |

---

**PHẦN IV. ĐIỀU CHỈNH SAU BÀI DẠY (NẾU CÓ)**
*(Đây là phần để trống để giáo viên ghi chú lại sau khi thực tế giảng dạy)*
1.  **Về nội dung, kiến thức:**
    * ......................................................................
2.  **Về phương pháp, kỹ thuật tổ chức:**
    * ......................................................................
3.  **Về học sinh (những khó khăn, điểm cần lưu ý):**
    * ......................................................................

---

**PHẦN V. PHIẾU BÀI TẬP (NẾU CÓ)**
(QUAN TRỌNG: Bạn CHỈ tạo phần này nếu DỮ LIỆU ĐẦU VÀO số 6 `{yeu_cau_phieu}` là 'CÓ'. Nếu là 'KHÔNG', hãy bỏ qua hoàn toàn phần này.)

- Nếu `{yeu_cau_phieu}` là 'CÓ':
- Hãy thiết kế một Phiếu bài tập (Worksheet) ngắn gọn, bám sát nội dung của **Hoạt động 3: Luyện tập / Thực hành**.
- Phiếu phải được trình bày sinh động, vui nhộn (dùng emojis 🌟, 🦋, 🖍️, 🐝, lời dẫn thân thiện).
- Đặt tên phiếu rõ ràng (ví dụ: PHIẾU BÀI TẬP - BÀI: {ten_bai}).
- Bao gồm 2-3 bài tập nhỏ (ví dụ: nối, điền từ, khoanh tròn).

---

**PHẦN VI. SƠ ĐỒ TƯ DUY (MÃ NGUỒN GRAPHVIZ)**
(QUAN TRỌNG: Bạn CHỈ tạo phần này nếu DỮ LIỆU ĐẦU VÀO số 7 `{yeu_cau_mindmap}` là 'CÓ'. Nếu là 'KHÔNG', hãy bỏ qua hoàn toàn phần này.)

- Nếu `{yeu_cau_mindmap}` là 'CÓ':
- **YÊU CẦU BẮT BUỘC:** Bạn PHẢI tạo một Sơ đồ tư duy (Mind Map) tóm tắt nội dung chính của bài học {ten_bai} bằng **ngôn ngữ Graphviz DOT**.
- **TUYỆT ĐỐI KHÔNG SỬ DỤNG:** Markdown, gạch đầu dòng, hay bất kỳ định dạng nào khác ngoài mã Graphviz DOT thuần túy trong phần này.
- **TUYỆT ĐỐI KHÔNG TẠO BẤT CỨ TIÊU ĐỀ NÀO** (ví dụ: PHẦN VI., hay bất kỳ dòng văn bản nào khác) **TRƯỚC THẺ START_GRAPHVIZ**.
- Sơ đồ phải rõ ràng, phân cấp, sử dụng tiếng Việt có dấu trong các nhãn (label) và **phải có nhãn mô tả ý tưởng chi tiết (để chức năng trích xuất gợi ý hoạt động hoạt động được)**. Sử dụng `layout=twopi` hoặc `layout=neato` để có bố cục tỏa tròn đẹp mắt.
- **QUAN TRỌNG:** Bọc toàn bộ mã code Graphviz DOT trong 2 thẻ **DUY NHẤT**: `[START_GRAPHVIZ]` ở dòng đầu tiên và `[END_GRAPHVIZ]` ở dòng cuối cùng của mã nguồn. Không thêm bất kỳ văn bản nào khác bên ngoài hai thẻ này trong phần VI.
- **YÊU CẦU MỚI VỀ SƠ ĐỒ:** Sơ đồ phải **đơn giản**, **trọng tâm**. Chỉ bao gồm 1 nút trung tâm (tên bài học) và 3-4 nhánh chính (các hoạt động/kiến thức cốt lõi). **Không rườm rà, ít chi tiết.**

---
Hãy bắt đầu tạo giáo án.
"""

# -----------------------------------------------------------------
# CÁC HÀM XỬ LÝ (GIỮ NGUYÊN XI TỪ FILE app (6).py CỦA BẠN)
# -----------------------------------------------------------------
def clean_content(text):
    # 1. Loại bỏ cụm "Cách tiến hành"
    text = re.sub(r'Cách tiến hành[:]*\s*', '', text, flags=re.IGNORECASE).strip()
    # 2. Loại bỏ TẤT CẢ các thẻ HTML (bao gồm <br>)
    text = re.sub(r'<[^>]+>', '', text, flags=re.IGNORECASE).strip()
    return text

def create_word_document(markdown_text, lesson_title):
    document = Document()
    
    # 1. Định nghĩa style (đã được tối ưu ở phiên bản trước)
    try:
        style_id = 1
        from docx.shared import Pt
             
        document.styles.add_style('ListBulletCustom', WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles['ListBulletCustom']
        style.base_style = document.styles['List Paragraph']
        font = style.font
        font.size = Pt(12)
    except Exception as e:
        pass
    
    if lesson_title:
        document.add_heading(f"KẾ HOẠCH BÀI DẠY: {lesson_title.upper()}", level=1)
        # Canh giữa tiêu đề
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()

    lines = markdown_text.split('\n')
    is_in_table_section = False
    is_in_part_vi = False 
    is_in_part_iii_section = False 
    table = None
    
    # --------------------------------------------------------------------------------
    # 2. TÁCH VÀ LƯU MÃ GRAPHVIZ
    # --------------------------------------------------------------------------------
    graph_code_content = ""
    parsing_graph = False
    
    for line in lines:
        if "[START_GRAPHVIZ]" in line:
            parsing_graph = True
            continue
        if "[END_GRAPHVIZ]" in line:
            parsing_graph = False
            continue 
        if parsing_graph:
            graph_code_content += line + "\n"
    # --------------------------------------------------------------------------------

    
    # Lặp qua toàn bộ nội dung gốc để tạo Word document
    for line in lines: 
        line = line.strip()
        if not line:
            continue
        # 🔹 Bỏ dòng tiêu đề "PHẦN VI. SƠ ĐỒ TƯ DUY (MÃ NGUỒN GRAPHVIZ)" nếu AI vẫn sinh ra
        if re.match(r'PHẦN\s*VI\.\s*SƠ\s*ĐỒ\s*TƯ\s*DUY', line, re.IGNORECASE):
            continue

        # *******************************************************************
        # BƯỚC 1: XỬ LÝ PHẦN VI (LOẠI BỎ CODE THÔ VÀ TIÊU ĐỀ THỪA)
        # *******************************************************************
        if re.match(r'PHẦN VI\.\s*SƠ ĐỒ TƯ DUY.*', line, re.IGNORECASE) or "[START_GRAPHVIZ]" in line:
            is_in_part_vi = True
            continue 
            
        if "[END_GRAPHVIZ]" in line:
            is_in_part_vi = False
            continue 
            
        if is_in_part_vi:
            # Loại bỏ mọi nội dung trong khi đang phân tích code Graphviz
            continue
            
        if line.startswith("PHẦN VI."):
             # Nếu AI vẫn tạo ra tiêu đề PHẦN VI. ngoài luồng, ta bỏ qua nó
             continue 
             
        # *******************************************************************
        
        # --------------------------------------------------------------------------------
        # BƯỚC 2: XỬ LÝ BẢNG CHÍNH (HOẠT ĐỘNG) - PHẦN III
        # --------------------------------------------------------------------------------
        # Bắt đầu bảng (Tiêu đề 2 cột)
        if re.match(r'\|.*Hoạt động của giáo viên.*\|.*Hoạt động của học sinh.*\|', line, re.IGNORECASE):
            is_in_table_section = True
            is_in_part_iii_section = False # Tắt cờ lọc nội dung thừa
            
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.autofit = False
            table.columns[0].width = Inches(3.5)
            table.columns[1].width = Inches(3.5)
            hdr_cells = table.rows[0].cells
            
            # Định dạng tiêu đề cột
            for cell in hdr_cells:
                cell.paragraphs[0].text = cell.paragraphs[0].text.strip()
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            hdr_cells[0].text = "Hoạt động của giáo viên"
            hdr_cells[1].text = "Hoạt động của học sinh"
            continue

        if is_in_table_section and table is not None:
            if line.startswith('| :---'):
                continue
                
            # Trường hợp 1: Dòng là một hàng hợp lệ trong bảng Markdown
            if line.startswith('|') and len(line.split('|')) >= 3:
                cells_content = [c.strip() for c in line.split('|')[1:-1]]

                if len(cells_content) == 2:
                    
                    gv_content = clean_content(cells_content[0].strip())
                    hs_content = clean_content(cells_content[1].strip())
                    
                    # QUAN TRỌNG: BỎ QUA DÒNG HƯỚNG DẪN MẪU (Ví dụ: (Viết chi tiết...))
                    if (gv_content.startswith('(') and gv_content.endswith(')')) and (hs_content.startswith('(') and hs_content.endswith(')')):
                        continue
                        
                    # Nếu nội dung của cả hai cột là rỗng (sau khi làm sạch), bỏ qua dòng này.
                    if not gv_content.strip() and not hs_content.strip():
                        continue
                    
                    # Regex để bắt các dòng Tiêu đề Hoạt động (chỉ bắt số thứ tự 1, 2, 3, 4)
                    ACTIVITY_HEADERS_PATTERN = re.compile(r'^\s*(\d+\.\sHoạt động.*)\s*', re.IGNORECASE | re.DOTALL)
                    is_main_header = ACTIVITY_HEADERS_PATTERN.match(gv_content)
                    
                    if is_main_header:
                        # TRƯỜNG HỢP: DÒNG LÀ TIÊU ĐỀ HOẠT ĐỘNG (CẦN MERGE CELL)
                        current_row = table.add_row().cells 
                        current_row[0].merge(current_row[1])
                        
                        # Xử lý nội dung 
                        content_line = gv_content.strip().strip('*-').strip()
                        if content_line:
                            p = current_row[0].add_paragraph(content_line)
                            p.runs[0].bold = True # In đậm Tiêu đề Chính
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    else:
                        # TRƯỜNG HỢP: DÒNG LÀ NỘI DUNG CHI TIẾT CỦA GV/HS (KHÔNG MERGE)
                        current_row = table.add_row().cells 

                        # Xử lý nội dung cho cột GV và HS
                        for cell_index, cell_content in enumerate([gv_content, hs_content]):
                            # Loại bỏ các dấu ** thừa trong nội dung chi tiết
                            cell_content = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_content)
                            
                            content_lines = cell_content.split('\n')
                            
                            for content_line in content_lines:
                                content_line = content_line.strip()
                                if not content_line: continue
                                
                                # Chỉ định dấu gạch đầu dòng (Sử dụng list bullet chuẩn)
                                if content_line.startswith('*') or content_line.startswith('-'):
                                    clean_text = content_line.lstrip('*- ').strip()
                                    # Tạo paragraph với style List Paragraph
                                    p = current_row[cell_index].add_paragraph(style='List Paragraph') 
                                    p.add_run('•\t') # Thêm dấu bullet thủ công
                                    p.add_run(clean_text)
                                    p.paragraph_format.left_indent = Inches(0.25)
                                else:
                                    current_row[cell_index].add_paragraph(content_line)
                    
                    continue # Chuyển sang dòng tiếp theo trong Markdown
            
            # Trường hợp 2 (FIX QUAN TRỌNG): Chỉ đóng bảng nếu gặp dấu phân cách '---' hoặc tiêu đề 'PHẦN IV.'
            if line == '---' or re.match(r'^PHẦN\sIV\.', line, re.IGNORECASE):
                is_in_table_section = False
                continue
                
            # Trường hợp 3: Nếu là dòng bất thường (text trôi nổi) nhưng chưa đến điểm kết thúc, BỎ QUA dòng đó.
            # Điều này giúp bảng vẫn tiếp tục nếu AI chèn text rác giữa HĐ 3 và HĐ 4.
            continue 
            
            # --------------------------------------------------------------------------------
            # BƯỚC 3: XỬ LÝ NỘI DUNG NGOÀI BẢNG (I, II, IV, V) + LỌC TEXT THỪA PHẦN III
            # --------------------------------------------------------------------------------
            
        # Bắt các tiêu đề chính (I, II, IV, V)
        if re.match(r'^[IVX]+\.\s|PHẦN\s[IVX]+\.', line):
            clean_line = line.strip().strip('**')
            document.add_heading(clean_line, level=2)
            
            # Bật cờ lọc nếu đây là Tiêu đề III. (Sau đó sẽ bị tắt khi tìm thấy bảng)
            if clean_line.startswith('III.'):
                is_in_part_iii_section = True
            else:
                is_in_part_iii_section = False
            
            continue

        # LỌC: Bỏ qua các dòng hướng dẫn/quy tắc thừa giữa tiêu đề III và bảng
        if is_in_part_iii_section and not is_in_table_section:
            continue
            
        # Các tiêu đề con (vd: 1. **Về kiến thức:**)
        elif line.startswith('**') and line.endswith('**'):
            document.add_heading(line.strip('**').replace('**', ''), level=3)

        # Danh sách gạch đầu dòng (List Bullet - Dấu chấm)
        elif line.startswith('*') or line.startswith('-'):
            clean_text = line.lstrip('*- ').strip().replace('**', '')
            p = document.add_paragraph(style='List Paragraph')
            p.add_run('•\t') 
            p.add_run(clean_text) 

            p.paragraph_format.left_indent = Inches(0.25)
        else:
            # Văn bản thường 
            document.add_paragraph(line.replace('**', ''))


    # *******************************************************************
    # 4. XỬ LÝ PHẦN VI (GỢI Ý SƠ ĐỒ TƯ DUY) 
    # *******************************************************************
    # Đảm bảo PHẦN VI. luôn được đặt ở cuối
    document.add_heading("PHẦN VI. GỢI Ý SƠ ĐỒ TƯ DUY", level=2)
                 
    if graph_code_content.strip():
        # Lọc các nhãn không cần thiết
        labels = re.findall(r'label="([^"]*)"', graph_code_content, re.DOTALL)
        unique_labels = sorted(list(set(label.strip() for label in labels if label.strip())))

        if unique_labels:
            document.add_paragraph("(Dưới đây là gợi ý nội dung chính (Key Ideas) được trích xuất từ sơ đồ tư duy do AI tạo. Giáo viên có thể dựa vào đây để vẽ hoặc chèn hình ảnh sơ đồ từ giao diện web.)")
            document.add_paragraph() 

            # Lọc nhãn trung tâm/chính (lấy nhãn dài nhất có chứa lesson_title)
            center_label = next((label for label in unique_labels if lesson_title.upper() in label.upper() and len(label) > 10), None)
            
            if center_label and center_label in unique_labels:
                
                center_label_parts = center_label.replace(r'\n', ' | ').split('|')
                
                # 1. Thêm nhãn trung tâm (Nhánh cấp 1 - Dùng dòng text đậm)
                p = document.add_paragraph(f"- {center_label_parts[0].replace('**', '').strip()}")
                p.runs[0].bold = True
                p.style = 'List Paragraph' 
                p.paragraph_format.left_indent = Inches(0.25)
                
                unique_labels.remove(center_label)
            
            # 2. Thêm các nhánh chính và nhánh phụ (Nhánh cấp 2, 3)
            for label in unique_labels:
                # Bỏ qua các nhãn ngắn, chỉ là tên node (như "center", "nhanh1")
                if (len(label) < 10 and not any(c.isalpha() for c in label)) or label.lower().strip() in ["center", "nhanh1", "nhanh2", "noidung", "nội dung", "kết quả", "cach_lam", "luyen_tap", "van_dung", "muc_tieu"]:
                    continue

                processed_label = label.replace(r'\n', '\n')
                label_parts = processed_label.split('\n')
                
                main_label = label_parts[0].strip().replace('**', '') 
                
                if not main_label or len(main_label) < 5: 
                    continue
                
                # Nhãn chính (main branch) - Cấp 2
                p = document.add_paragraph(style='List Paragraph')
                p.add_run('•\t')
                p.add_run(f"  {main_label}")
                p.paragraph_format.left_indent = Inches(0.5)
                        
                # Thêm các dòng phụ (sub branch) - Cấp 3
                for part in label_parts[1:]:
                    part = part.strip().replace('**', '') 
                    if part and len(part) > 3: 
                        p = document.add_paragraph(style='List Paragraph')
                        p.add_run('•\t')
                        p.add_run(f"    {part}")
                        p.paragraph_format.left_indent = Inches(0.75)

        else:
            document.add_paragraph("(AI đã tạo Graphviz nhưng không tìm thấy nhãn nội dung (label) chi tiết để trích xuất gợi ý. Vui lòng kiểm tra lại yêu cầu Sơ đồ tư duy.)")
        
    else:
        document.add_paragraph("(Không tìm thấy mã nguồn Graphviz. Có thể yêu cầu tạo sơ đồ tư duy là 'KHÔNG' hoặc AI đã không tạo ra nội dung Graphviz hợp lệ.)")
    
    # *******************************************************************
    
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio


# -----------------------------------------------------------------
# 5. XÂY DỰNG GIAO DIỆN "CHAT BOX" (Web App) (GIỮ NGUYÊN)
# -----------------------------------------------------------------

st.set_page_config(page_title="Trợ lý Soạn giáo án AI", page_icon="🤖")
st.title("🤖 Trợ lý Soạn Kế hoạch bài dạy Tiểu học")
st.write("Sản phẩm của thầy giáo Hoàng Trọng Nghĩa.")


st.markdown("*(Kế hoạch bài dạy được biên soạn theo chuẩn Chương trình GDPT 2018)*")


# Tạo 5 ô nhập liệu cho 5 biến số
mon_hoc = st.text_input("1. Môn học:", placeholder="Ví dụ: Tiếng Việt")
lop = st.text_input("2. Lớp:", placeholder="Ví dụ: 2")
bo_sach = st.text_input("3. Bộ sách:", placeholder="Ví dụ: Cánh Diều")
ten_bai = st.text_input("4. Tên bài học / Chủ đề:", placeholder="Ví dụ: Bài 2: Thời gian của em")
yeu_cau = st.text_area("5. Yêu cầu cần đạt:", placeholder="Điền Yêu cầu cần đạt ...", height=150)

# -----------------------------------------------------------------
# 6. KHAI BÁO BIẾN CHO FILE UPLOADER (GIỮ NGUYÊN LOGIC TẢI 2 ẢNH)
# -----------------------------------------------------------------
uploaded_files = st.file_uploader( # Đổi tên biến (từ _file sang _files)
    "6. [Tải Lên] Ảnh/PDF trang Bài tập SGK (Tối đa 2 ảnh, Tùy chọn)", # Sửa label
    type=["pdf", "png", "jpg", "jpeg"],
    accept_multiple_files=True # CHO PHÉP TẢI NHIỀU FILE
)
# -----------------------------------------------------------------


# 7. KHAI BÁO BIẾN CHO CHECKBOX PHIẾU BÀI TẬP (GIỮ NGUYÊN)
tao_phieu = st.checkbox("7. Yêu cầu tạo kèm Phiếu Bài Tập", value=False)

# 8. <-- MỚI: Thêm Checkbox cho Sơ đồ tư duy (GIỮ NGUYÊN)
tao_mindmap = st.checkbox("8. Yêu cầu tạo Sơ đồ tư duy trực quan", value=True)

# Nút bấm để tạo giáo án
if st.button("🚀 Tạo KH bài dạy ngay!"):
    if not mon_hoc or not lop or not bo_sach or not ten_bai or not yeu_cau:
        st.error("Vui lòng nhập đầy đủ cả 5 thông tin!")
    else:
        with st.spinner("Trợ lý AI đang soạn giáo án, vui lòng chờ trong giây lát..."):
            try:
                # Logic cho Biến số Tùy chọn 1 (Tạo Phiếu Bài Tập) (GIỮ NGUYÊN)
                yeu_cau_phieu_value = "CÓ" if tao_phieu else "KHÔNG"

                # Logic cho Biến số Tùy chọn 2 (Sơ đồ tư duy) (GIỮ NGUYÊN)
                yeu_cau_mindmap_value = "CÓ" if tao_mindmap else "KHÔNG"


                # 1. Chuẩn bị Nội dung (Content List) cho AI (Tích hợp File và Text)
                content = []

                # 2. Điền Prompt (7 biến số text) (GIỮ NGUYÊN)
                final_prompt = PROMPT_GOC.format(
                    mon_hoc=mon_hoc,
                    lop=lop,
                    bo_sach=bo_sach,
                    ten_bai=ten_bai,
                    yeu_cau=yeu_cau,
                    yeu_cau_phieu=yeu_cau_phieu_value,
                    yeu_cau_mindmap=yeu_cau_mindmap_value
                )

                # -----------------------------------------------------------------
                # 3. LOGIC XỬ LÝ ẢNH (GIỮ NGUYÊN LOGIC XỬ LÝ 2 ẢNH)
                # -----------------------------------------------------------------
                if uploaded_files: 
                    files_to_process = uploaded_files[:2]
                    
                    st.info(f"Đang phân tích {len(files_to_process)} ảnh bài tập...")

                    for uploaded_file in files_to_process:
                        if uploaded_file.type == "application/pdf":
                            st.error(f"Lỗi: File {uploaded_file.name} là PDF, chưa được hỗ trợ. Vui lòng tải file ảnh (PNG, JPG).")
                            continue 
                        try:
                            image = Image.open(uploaded_file)
                            content.append(image)
                        except Exception as e:
                            st.error(f"Lỗi khi mở ảnh {uploaded_file.name}: {e}")
                # -----------------------------------------------------------------


                # 4. Thêm Prompt vào danh sách Content (luôn luôn có)
                content.append(final_prompt)

                # 5. Gọi AI với danh sách nội dung (content)
                response = model.generate_content(content)

                # 6. Hiển thị kết quả (GIỮ NGUYÊN)
                st.balloons()
                st.subheader("🎉 Giáo án của bạn đã sẵn sàng:")

                # LÀM SẠCH KẾT QUẢ ĐỂ CHỈ HIỂN THỊ GIÁO ÁN
                full_text = response.text

                # Lọc sạch thẻ <br> (lỗi cũ)
                full_text = re.sub(r'<\s*br\s*\/?>', '\n', full_text, flags=re.IGNORECASE)

                start_index = full_text.find("I. Yêu cầu cần đạt")

                if start_index != -1:
                    cleaned_text = full_text[start_index:]
                else:
                    cleaned_text = full_text

                # --- BỘ LỌC TOÀN DIỆN CHO HIỂN THỊ WEB ---
                cleaned_text_display = cleaned_text
                
                # 1. LỌC CỤM "Cách tiến hành:"
                cleaned_text_display = re.sub(r'Cách tiến hành[:]*\s*', '', cleaned_text_display, flags=re.IGNORECASE)
                
                # 2. LỌC CÁC DÒNG QUY TẮC/HƯỚNG DẪN THỪA TRONG PHẦN III
                cleaned_text_display = re.sub(r'\*\*QUY TẮC QUAN TRỌNG.*', '', cleaned_text_display, flags=re.IGNORECASE)
                
                # 3. LỌC TIÊU ĐỀ GRAPHVIZ THÔ
                cleaned_text_display = re.sub(r'PHẦN VI\.\s*SƠ ĐỒ TƯ DUY.*', '', cleaned_text_display, flags=re.IGNORECASE)
                # ---------------------------------------------------

                # --- KHỐI LOGIC HIỂN THỊ SƠ ĐỒ TƯ DUY TRÊN WEB (GIỮ NGUYÊN) ---
                start_tag = "[START_GRAPHVIZ]"
                end_tag = "[END_GRAPHVIZ]"

                # (Sửa lỗi: dùng full_text gốc để tìm sơ đồ, vì cleaned_text_display đã bị xóa mất PHẦN VI)
                if tao_mindmap and start_tag in full_text:
                    try:
                        # Tách nội dung trước và sau code Graphviz
                        before_graph = cleaned_text_display.split(start_tag)[0]
                        
                        # Tách code từ full_text gốc
                        temp = full_text.split(start_tag)[1]
                        graph_code = temp.split(end_tag)[0].strip()
                        
                        # Tách nội dung sau code
                        after_graph = cleaned_text_display.split(end_tag)[1]

                        # Hiển thị
                        st.markdown(before_graph)
                        st.subheader("Sơ đồ tư duy (Mind Map) - VẼ TRỰC TIẾP:")
                        if graph_code:
                            st.graphviz_chart(graph_code) # Vẽ sơ đồ
                        else:
                            st.warning("AI đã tạo thẻ tag nhưng mã nguồn Graphviz rỗng. Vui lòng chạy lại.")
                        
                        # Loại bỏ tiêu đề "PHẦN VI." nếu nó nằm trong `after_graph`
                        after_graph = re.sub(r'PHẦN VI\.\s*GỢI Ý SƠ ĐỒ TƯ DUY.*', '', after_graph, flags=re.IGNORECASE)
                        st.markdown(after_graph)

                    except IndexError:
                        st.error("Lỗi khi trích xuất mã nguồn Graphviz: Không tìm thấy thẻ đóng `[END_GRAPHVIZ]`.")
                        st.markdown(cleaned_text_display) # Hiển thị văn bản (đã bị xóa code)
                    except Exception as e:
                        st.error(f"Lỗi khi vẽ sơ đồ tư duy: {e}")
                        st.markdown(cleaned_text_display)
                else:
                    st.markdown(cleaned_text_display) # Hiển thị nếu không tick chọn hoặc không có code
                # --- KẾT THÚC KHỐI LOGIC SƠ ĐỒ TƯ DU DUY ---


                # BẮT ĐẦU KHỐI CODE TẢI XUỐNG WORD
                # Hàm create_word_document đã được cập nhật để loại bỏ nội dung thừa/mã thô
                word_bytes = create_word_document(cleaned_text, ten_bai)


                st.download_button(
                    label="⬇️ Tải về Giáo án (Word)",
                    data=word_bytes,
                    file_name=f"GA_{ten_bai.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                # Xử lý lỗi đặc biệt khi API Key bị lỗi (chỉ cần một dòng thông báo)
                if "API_KEY" in str(e):
                    st.error("Lỗi xác thực API: Vui lòng kiểm tra lại 'GEMINI_API_KEY' trong Streamlit Secrets.")
                else:
                    st.error(f"Đã có lỗi xảy ra: {e}")
                    st.error("Lỗi này có thể do AI không tạo ra đúng định dạng hoặc có lỗi kết nối.")

# BẮT ĐẦU PHẦN SIDEBAR
st.sidebar.title("Giới thiệu")
st.sidebar.info(
"""
Sản phẩm của thầy giáo Hoàng Trọng Nghĩa, Trường Tiểu học Hồng Gai - Tỉnh Quảng Ninh. tham gia ngày hội "Nhà giáo sáng tạo với công nghệ số và trí tuệ nhân tạo".

Sản phẩm ứng dụng AI để tự động soạn Kế hoạch bài dạy cho giáo viên Tiểu học theo đúng chuẩn Chương trình GDPT 2018.
"""
)
